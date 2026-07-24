"""위키 노트(마크다운) → Word(.docx) 내보내기 (Phase 2, 2026-07-24).

옵시디언을 쓰지 않는 사용자를 위해, 허브 노트와 동일한 내용을 편집 가능한
Word 문서로 저장한다. python-docx만 사용(순수 파이썬·시스템 의존성 없음)."""
from __future__ import annotations

import re
import tempfile
import shutil
from pathlib import Path


def _safe_name(stem: str) -> str:
    return re.sub(r'[/\\:*?"<>|]', "_", stem).strip() or "wiki"


def _add_inline(paragraph, text: str) -> None:
    """**굵게** 정도만 처리한 인라인 런 추가."""
    for i, seg in enumerate(re.split(r"(\*\*[^*]+\*\*)", text)):
        if not seg:
            continue
        if seg.startswith("**") and seg.endswith("**"):
            paragraph.add_run(seg[2:-2]).bold = True
        else:
            paragraph.add_run(seg)


def note_md_to_docx(md: str, out_path: Path, *, meta: dict | None = None) -> Path:
    """마크다운 노트 문자열을 .docx로 변환해 out_path에 저장."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # ── frontmatter 분리 → 제목·서지 헤더 ──
    body = md
    fm: dict[str, str] = {}
    m = re.match(r"^---\n(.*?)\n---\n?(.*)$", md, re.S)
    if m:
        for line in m.group(1).splitlines():
            mm = re.match(r"^(\w+):\s*(.*)$", line)
            if mm:
                fm[mm.group(1)] = mm.group(2).strip().strip('"')
        body = m.group(2)
    fm.update(meta or {})

    title = fm.get("title") or (meta or {}).get("title") or ""
    if title:
        doc.add_heading(title, level=0)
    _bib = " · ".join(
        x for x in [fm.get("author", ""), fm.get("published", ""), fm.get("publisher", "")] if x
    )
    if _bib:
        _p = doc.add_paragraph()
        _r = _p.add_run(_bib)
        _r.italic = True
        _r.font.size = Pt(10)

    # ── 본문 라인 단위 변환 ──
    lines = body.splitlines()
    i = 0
    while i < len(lines):
        ln = lines[i].rstrip()
        if not ln.strip():
            i += 1
            continue
        # 표(| ... |) — 연속 파이프 줄 묶기 (구분선 |---| 은 건너뜀)
        if ln.lstrip().startswith("|"):
            rows = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not re.match(r"^\s*:?-{2,}", cells[0] if cells else ""):
                    rows.append(cells)
                i += 1
            if rows:
                ncol = max(len(r) for r in rows)
                tbl = doc.add_table(rows=0, cols=ncol)
                tbl.style = "Light Grid Accent 1"
                for r in rows:
                    cells = tbl.add_row().cells
                    for c in range(ncol):
                        cells[c].text = r[c] if c < len(r) else ""
            continue
        # 헤딩
        h = re.match(r"^(#{1,6})\s+(.*)$", ln)
        if h:
            level = min(len(h.group(1)), 4)
            doc.add_heading(h.group(2), level=level)
            i += 1
            continue
        # 인용
        if ln.lstrip().startswith(">"):
            _p = doc.add_paragraph(style="Intense Quote")
            _add_inline(_p, ln.lstrip()[1:].strip())
            i += 1
            continue
        # 키워드 해시태그 줄 (#키워드 — 해설)
        if ln.lstrip().startswith("#") and not ln.lstrip().startswith("##"):
            _p = doc.add_paragraph()
            _p.add_run(ln.strip()).bold = True
            i += 1
            continue
        # 불릿
        if re.match(r"^\s*[-*]\s+", ln):
            _p = doc.add_paragraph(style="List Bullet")
            _add_inline(_p, re.sub(r"^\s*[-*]\s+", "", ln))
            i += 1
            continue
        # 일반 문단
        _p = doc.add_paragraph()
        _add_inline(_p, ln)
        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def build_docx_from_chapter_summaries(ws_name: str, stem: str, out_dir: Path) -> tuple[bool, str]:
    """챕터 요약 → 허브 노트(임시 생성) → .docx. (ok, path or msg)."""
    from services.wiki import build_wiki_from_chapter_summaries
    tmp = Path(tempfile.mkdtemp(prefix="mb_docx_"))
    try:
        ok, msg = build_wiki_from_chapter_summaries(ws_name, stem, wiki_dir=tmp)
        if not ok:
            return False, msg
        md_path = Path(msg)
        if not md_path.exists():
            return False, "노트 생성 실패"
        md = md_path.read_text(encoding="utf-8")
        out_path = out_dir / (_safe_name(stem) + ".docx")
        note_md_to_docx(md, out_path)
        return True, str(out_path)
    except Exception as e:
        return False, f"{type(e).__name__}: {str(e)[:150]}"
    finally:
        shutil.rmtree(tmp, ignore_errors=True)
